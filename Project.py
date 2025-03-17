import os
import pandas as pd
import numpy as np
from docx import Document
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
import lightgbm as lgb

# Папки с файлами
correct_folder = "good"
incorrect_folder = "bad"

def extract_features(doc_path):
    try:
        doc = Document(doc_path)
        first_paragraph = doc.paragraphs[0] if doc.paragraphs else None

        # Извлечение базовых параметров
        font_name = "Unknown"
        font_size = np.nan
        alignment = "Unknown"
        
        if first_paragraph and first_paragraph.runs:
            font_name = first_paragraph.runs[0].font.name or "Unknown"
            font_size = first_paragraph.runs[0].font.size.pt if first_paragraph.runs[0].font.size else np.nan
            alignment = first_paragraph.alignment if first_paragraph.alignment else "Unknown"

        # Средний отступ в документе
        indents = [p.paragraph_format.first_line_indent.pt if p.paragraph_format.first_line_indent else 0 for p in doc.paragraphs]
        avg_indent = np.mean(indents) if indents else 0

        # Количество заголовков
        num_headings = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))

        return font_name, font_size, alignment, avg_indent, num_headings
    except Exception as e:
        print(f"Ошибка обработки {doc_path}: {e}")
        return "Unknown", np.nan, "Unknown", np.nan, np.nan


def load_data():
    """ Загружает данные из папок и формирует DataFrame """
    data = []
    
    for folder, label in [(correct_folder, 1), (incorrect_folder, 0)]:
        for file in os.listdir(folder):
            if file.endswith(".docx"):
                file_path = os.path.join(folder, file)
                font_name, font_size, alignment, avg_indent, num_headings = extract_features(file_path)
                data.append([file, font_name, font_size, alignment, avg_indent, num_headings, label])

    df = pd.DataFrame(data, columns=["file_name", "font_name", "font_size", "alignment", "avg_indent", "num_headings", "label"])
    
    # Преобразуем alignment в строки
    df["alignment"] = df["alignment"].astype(str)

    # Заполняем пропущенные значения корректным методом
    df = df.fillna({
        "font_size": df["font_size"].median(),
        "avg_indent": df["avg_indent"].median(),
        "num_headings": 0
    })
    
    # Кодирование категориальных данных
    le_font = LabelEncoder()
    df["font_name"] = le_font.fit_transform(df["font_name"])
    
    le_align = LabelEncoder()
    df["alignment"] = le_align.fit_transform(df["alignment"])
    
    return df, le_font, le_align


def train_model(df):
    """ Обучает модель на загруженных данных """
    X = df.drop(columns=["file_name", "label"])
    y = df["label"]

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    model = lgb.LGBMClassifier()
    model.fit(X_train, y_train)

    accuracy = model.score(X_test, y_test)
    print(f"Точность модели: {accuracy:.2f}")

    return model

def predict_new_file(model, file_path, le_font, le_align):
    """ Проверяет новый файл на соответствие ГОСТу """
    font_name, font_size, alignment, avg_indent, num_headings = extract_features(file_path)
    
    # Обработка неизвестных значений
    font_name = font_name if font_name in le_font.classes_ else "Unknown"
    alignment = alignment if alignment in le_align.classes_ else "Unknown"
    
    font_name_encoded = le_font.transform([font_name])[0]
    alignment_encoded = le_align.transform([alignment])[0]
    
    # Формируем DataFrame для предсказания
    features = pd.DataFrame([[font_name_encoded, font_size, alignment_encoded, avg_indent, num_headings]], 
                            columns=["font_name", "font_size", "alignment", "avg_indent", "num_headings"])

    prediction = model.predict(features)[0]
    result = "Соответствует ГОСТу" if prediction == 1 else "НЕ соответствует ГОСТу"
    print(f"Файл {file_path}: {result}")
    return result

# Запуск
if __name__ == "__main__":
    df, le_font, le_align = load_data()
    df.to_csv("dataset.csv", index=False)
    print(df.head())

    model = train_model(df)
    
    test_file = "TEST.docx"  # Укажи путь к тестовому файлу
    if os.path.exists(test_file):
        predict_new_file(model, test_file, le_font, le_align)
    else:
        print(f"Файл {test_file} не найден! Добавьте тестовый документ.")
